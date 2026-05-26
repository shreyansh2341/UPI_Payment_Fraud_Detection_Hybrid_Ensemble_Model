import docx
import os

def replace_text_in_doc(doc, old_text, new_text):
    """Replaces text in paragraphs and tables within a docx document."""
    replacements_made = 0
    
    # Replace in paragraphs
    for para in doc.paragraphs:
        if old_text in para.text:
            # We want to preserve formatting, so we iterate through runs
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    replacements_made += 1
            # If the text spans multiple runs, fallback to full text replace (loses run-specific formatting)
            if old_text in para.text and old_text not in [r.text for r in para.runs]:
                para.text = para.text.replace(old_text, new_text)
                replacements_made += 1

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old_text in para.text:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                replacements_made += 1
                        if old_text in para.text and old_text not in [r.text for r in para.runs]:
                            para.text = para.text.replace(old_text, new_text)
                            replacements_made += 1
                            
    return replacements_made

def update_report(filepath, output_path):
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        return
        
    doc = docx.Document(filepath)
    print("Document loaded. Performing text replacements...")
    
    # Define replacements based on user feedback and actual model (Hybrid Ensemble: AE + XGB/RF)
    # Avoiding "v3" as requested.
    replacements = {
        # General Architecture
        "Convolutional Neural Networks (CNNs) are applied for feature extraction, Long Short-Term Memory networks (LSTMs) for analysing sequential transaction patterns, and Autoencoders for detecting anomalies": "A Hybrid Deep Learning and Machine Learning Model is applied, utilizing an Autoencoder for anomaly detection and feature engineering, combined with an XGBoost and Random Forest ensemble for precise supervised classification",
        "CNN, LSTM, and Autoencoder": "Autoencoder, XGBoost, and Random Forest",
        "CNN, LSTM, and Autoencoders": "Autoencoder, XGBoost, and Random Forest",
        "CNN, LSTM, and Autoencoder architectures": "Autoencoder and tree-based ensemble architectures",
        "CNN for feature extraction, LSTM for understanding transaction sequences, and Autoencoders for anomaly detection": "an Autoencoder for unsupervised anomaly detection and an XGBoost and Random Forest ensemble for supervised classification",
        "CNN-LSTM": "Autoencoder-XGBoost",
        "CNN\u2013LSTM": "Autoencoder-XGBoost",
        "RNNs and LSTMs for sequential analysis": "an XGBoost and Random Forest ensemble for high-precision classification",
        "Convolutional Neural Networks (CNNs)": "Machine Learning Classifiers",
        "Long Short-Term Memory networks (LSTMs)": "Tree-based Ensembles (XGBoost & Random Forest)",
        
        # Specific mentions
        "CNN-LSTM path": "XGBoost and Random Forest path",
        
        # Typos / Grammar cleanup observed in original
        "fuelled scams": "fueled scams",
        "analysing": "analyzing",
        "behaviors": "behaviours",
    }
    
    total_replacements = 0
    for old, new in replacements.items():
        count = replace_text_in_doc(doc, old, new)
        total_replacements += count
        print(f"Replaced {count} instances of: '{old[:30]}...' -> '{new[:30]}...'")
        
    doc.save(output_path)
    print(f"\nSaved updated report to: {output_path}")
    print(f"Total replacements made: {total_replacements}")

if __name__ == "__main__":
    input_report = r"c:\Users\Shreyansh Rai\OneDrive\Desktop\Fraud_Detection_Model_Paysim\Fraud_Detection_Model_Paysim_CC\PRECISE AND ENHANCEMENT OF UPI BASED TRANSACTION SCAM IDENTIFICATION USING HYBRID DL MODEL.docx"
    
    base_dir = os.path.dirname(input_report)
    output_report = os.path.join(base_dir, "PRECISE_AND_ENHANCEMENT_OF_UPI_BASED_TRANSACTION_SCAM_IDENTIFICATION_UPDATED.docx")
    
    update_report(input_report, output_report)
