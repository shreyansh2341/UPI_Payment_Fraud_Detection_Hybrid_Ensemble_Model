import os
import base64
import zlib
import urllib.request
import docx
from docx.shared import Inches

def generate_mermaid_via_kroki(mermaid_code, output_path):
    print(f"Generating diagram: {output_path}")
    compressed = zlib.compress(mermaid_code.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
    url = f"https://kroki.io/mermaid/png/{encoded}"
    
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    try:
        with urllib.request.urlopen(req) as response:
            with open(output_path, 'wb') as f:
                f.write(response.read())
    except Exception as e:
        print(f"Failed to generate {output_path}: {e}")

def create_system_architecture(out_dir):
    code = """
    graph TB
        subgraph Input
            Raw[Raw Transaction Data] --> Pre(Data Validation & Cleansing)
        end
        subgraph Feature Engineering
            Pre --> FE(18 Temporal & Behavioral Features)
        end
        subgraph Two-Path Model Evaluation
            FE --> PathB((Path B: Unsupervised<br>Autoencoder))
            FE --> PathA((Path A: Supervised<br>XGBoost + Random Forest))
            PathB -.->|Reconstruction Error 19th Feature| PathA
        end
        subgraph Output Logic
            PathA --> Dec{Is Prob >= 0.77?}
            PathB --> Anom{Is Anomaly Score High?}
            Dec -->|Yes| B[BLOCK: Confirmed Fraud]
            Dec -->|No| Anom
            Anom -->|Yes| R[REVIEW: Novel Fraud Alert]
            Anom -->|No| A[ALLOW: Legitimate]
        end
        style B fill:#ffcccc
        style R fill:#ffeecc
        style A fill:#ccffcc
    """
    generate_mermaid_via_kroki(code, os.path.join(out_dir, "system_architecture.png"))

def create_uml_diagram(out_dir):
    code = """
    classDiagram
        class TransactionSystem {
            +processTransaction(data)
            +generateFeatures(data)
        }
        class Autoencoder {
            +getReconstructionError(features) float
            +flagAnomaly(error) boolean
        }
        class SupervisedEnsemble {
            +predictFraudProb(features, aeError) float
            +classify(prob) string
        }
        class DecisionEngine {
            +evaluatePaths(aeResult, ensembleResult) string
            +blockTransaction()
            +flagForReview()
            +allowTransaction()
        }
        TransactionSystem --> Autoencoder
        TransactionSystem --> SupervisedEnsemble
        Autoencoder --> SupervisedEnsemble : provides aeError
        Autoencoder --> DecisionEngine
        SupervisedEnsemble --> DecisionEngine
    """
    generate_mermaid_via_kroki(code, os.path.join(out_dir, "uml_diagram.png"))

def embed_images_in_docx(doc_path, image_map, out_path):
    # image_map is dict mapping caption strings to image paths
    print(f"Loading DOCX: {doc_path}")
    doc = docx.Document(doc_path)
    
    replacements = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        matched_caption = None
        for caption, img_path in image_map.items():
            if text.startswith(caption) or caption in text:
                matched_caption = caption
                break
                
        if matched_caption:
            img_path = image_map[matched_caption]
            if os.path.exists(img_path):
                print(f"Found caption '{matched_caption}'. Inserting image {img_path}")
                # Try to remove old image in previous 3 paragraphs
                for prev_idx in range(max(0, i-3), i):
                    prev_p = doc.paragraphs[prev_idx]
                    for run in prev_p.runs:
                        if "graphic" in run._element.xml:
                            run.clear()
                            print(f" -> Removed old image at paragraph {prev_idx}")
                
                # Insert the new image right before the caption paragraph
                new_p = para.insert_paragraph_before()
                new_p.alignment = 1 # Center align
                run = new_p.add_run()
                # Use a specific width
                try:
                    run.add_picture(img_path, width=Inches(5.5))
                    replacements += 1
                    print(f" -> Embedded {img_path} successfully")
                except Exception as e:
                    print(f" -> Failed to embed {img_path}: {e}")
            else:
                print(f"Image not found on disk: {img_path}")
                
    doc.save(out_path)
    print(f"Saved modified DOCX to {out_path} with {replacements} new images embedded.")

if __name__ == "__main__":
    out_dir = r"c:\Users\Shreyansh Rai\OneDrive\Desktop\Fraud_Detection_Model_Paysim\Fraud_Detection_Model_Paysim_CC\docs\dig_report"
    os.makedirs(out_dir, exist_ok=True)
    
    print("Generating additional Architecture and UML diagrams...")
    create_system_architecture(out_dir)
    create_uml_diagram(out_dir)
    
    # We map the exact caption strings seen in report_content.json to the PNG files
    # 1. "Fig 5.1 System Architecture"
    # 2. "Fig 5.2 Flow Chart"
    # 3. "Fig.7.1 UML Diagram" (Ah, in chapter 7 it says 7.1 UML Diagram)
    # 4. "Fig. 7.2 DFD Level 0"
    # 5. "Fig. 7.3 DFD Level 1"
    # 6. "Fig. 7.4 DFD Level 2"
    # 7. "Fig 7.5 Use Case"
    
    image_mapping = {
        "Fig 5.1 System Architecture": os.path.join(out_dir, "system_architecture.png"),
        "Fig 5.2 Flow Chart": os.path.join(out_dir, "flowchart.png"),
        "Fig.7.1 UML Diagram": os.path.join(out_dir, "uml_diagram.png"),
        "Fig. 7.2 DFD Level 0": os.path.join(out_dir, "dfd_level_0.png"),
        "Fig. 7.3 DFD Level 1": os.path.join(out_dir, "dfd_level_1.png"),
        "Fig. 7.4 DFD Level 2": os.path.join(out_dir, "dfd_level_2.png"),
        "Fig 7.5 Use Case": os.path.join(out_dir, "use_case_diagram.png")
    }
    
    # We will use the already-text-updated docx as the base
    base_docx = r"c:\Users\Shreyansh Rai\OneDrive\Desktop\Fraud_Detection_Model_Paysim\Fraud_Detection_Model_Paysim_CC\PRECISE_AND_ENHANCEMENT_OF_UPI_BASED_TRANSACTION_SCAM_IDENTIFICATION_UPDATED.docx"
    
    # Let's overwrite the ORIGINAL requested file so the user actually sees the physical changes right where they asked.
    final_docx = r"c:\Users\Shreyansh Rai\OneDrive\Desktop\Fraud_Detection_Model_Paysim\Fraud_Detection_Model_Paysim_CC\PRECISE AND ENHANCEMENT OF UPI BASED TRANSACTION SCAM IDENTIFICATION USING HYBRID DL MODEL.docx"
    
    # If the base docx wasn't successfully created earlier, fall back to the original docx
    if not os.path.exists(base_docx):
        base_docx = final_docx
        print("Using original docx as base...")
    
    print("Embedding diagrams into the report DOCX...")
    embed_images_in_docx(base_docx, image_mapping, final_docx)
