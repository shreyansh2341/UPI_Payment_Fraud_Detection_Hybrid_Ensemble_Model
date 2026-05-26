import docx
import os
import json

def extract_docx_content(filepath):
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        return
    
    doc = docx.Document(filepath)
    content = {
        "paragraphs": [],
        "tables": [],
        "headings": []
    }
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            content["paragraphs"].append({
                "index": i,
                "text": para.text,
                "style": para.style.name
            })
            if para.style.name.startswith("Heading"):
                content["headings"].append({
                    "index": i,
                    "text": para.text,
                    "level": para.style.name
                })
                
    for i, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        content["tables"].append({
            "index": i,
            "data": table_data
        })
        
    output_path = "report_content.json"
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(content, f, indent=2)
    print(f"Extracted content saved to {output_path}")
    print(f"Total paragraphs: {len(content['paragraphs'])}")
    print(f"Total headings: {len(content['headings'])}")
    print(f"Total tables: {len(content['tables'])}")

if __name__ == "__main__":
    report_path = r"c:\Users\Shreyansh Rai\OneDrive\Desktop\Fraud_Detection_Model_Paysim\Fraud_Detection_Model_Paysim_CC\PRECISE AND ENHANCEMENT OF UPI BASED TRANSACTION SCAM IDENTIFICATION USING HYBRID DL MODEL.docx"
    extract_docx_content(report_path)
